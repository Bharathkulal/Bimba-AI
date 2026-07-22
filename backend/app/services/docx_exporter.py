import io
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_docx_resume(resume_data: dict) -> io.BytesIO:
    """
    Generates a professional DOCX resume document using python-docx.
    resume_data: dict containing keys:
        - master: dict (name, phone, address, linkedin, github, portfolio, website, summary, career_objective, etc.)
        - education: list of dicts (institution, degree, board, percentage, cgpa, passing_year, achievements)
        - experience: list of dicts (company, position, duration, description)
        - projects: list of dicts (name, description, tech_stack, role, duration)
        - skills: list of dicts (category, name, level)
        - certificates: list of dicts (name, organization, issue_date)
        - achievements_list or achievements: dict (hackathons, awards, soft_skills, extracurricular)
        - student: dict (optional fallback info)
    """
    doc = Document()
    
    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    # Helper to set colors
    COLOR_PRIMARY = RGBColor(30, 58, 138)  # Navy #1E3A8A
    COLOR_TEXT = RGBColor(31, 41, 55)      # Charcoal #1F2937
    COLOR_MUTED = RGBColor(75, 85, 99)     # Gray #4B5563
    
    master = resume_data.get("master", {})
    student = resume_data.get("student", {})
    
    # Use master values or fallbacks
    name = master.get("name") or student.get("student_name") or "Your Name"
    email = student.get("email") or "email@example.com"
    phone = master.get("phone") or ""
    address = master.get("address") or ""
    
    linkedin = master.get("linkedin") or ""
    github = master.get("github") or ""
    portfolio = master.get("portfolio") or master.get("website") or ""
    
    # 1. Header (Centered Name & Contact)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = title_p.add_run(name)
    run_name.font.name = 'Arial'
    run_name.font.size = Pt(18)
    run_name.font.bold = True
    run_name.font.color.rgb = COLOR_PRIMARY
    
    contact_parts = []
    if email: contact_parts.append(email)
    if phone: contact_parts.append(phone)
    if address: contact_parts.append(address)
    
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_p.add_run("  •  ".join(contact_parts))
    contact_run.font.name = 'Arial'
    contact_run.font.size = Pt(9.5)
    contact_run.font.color.rgb = COLOR_MUTED
    
    socials = []
    if linkedin: socials.append(f"LinkedIn: {linkedin}")
    if github: socials.append(f"GitHub: {github}")
    if portfolio: socials.append(f"Portfolio: {portfolio}")
    
    if socials:
        socials_p = doc.add_paragraph()
        socials_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        socials_run = socials_p.add_run("  |  ".join(socials))
        socials_run.font.name = 'Arial'
        socials_run.font.size = Pt(9)
        socials_run.font.color.rgb = COLOR_MUTED
        
    doc.add_paragraph() # Spacer
    
    def add_section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text.upper())
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        
        # Add thin border/divider line under heading
        p_border = doc.add_paragraph()
        p_border.paragraph_format.space_after = Pt(6)
        r_border = p_border.add_run("─" * 70)
        r_border.font.color.rgb = COLOR_MUTED
        r_border.font.size = Pt(6)

    # 2. Professional Summary
    summary = master.get("summary") or master.get("career_objective")
    if summary:
        add_section_heading("Professional Summary")
        sum_p = doc.add_paragraph()
        sum_run = sum_p.add_run(summary)
        sum_run.font.name = 'Arial'
        sum_run.font.size = Pt(10)
        sum_run.font.color.rgb = COLOR_TEXT
        
    # 3. Education
    education = resume_data.get("education", [])
    if education:
        add_section_heading("Education")
        for edu in education:
            p = doc.add_paragraph()
            inst = p.add_run(f"{edu.get('institution')} — ")
            inst.font.bold = True
            inst.font.name = 'Arial'
            inst.font.size = Pt(10)
            inst.font.color.rgb = COLOR_TEXT
            
            deg = p.add_run(edu.get('degree'))
            deg.font.name = 'Arial'
            deg.font.size = Pt(10)
            deg.font.color.rgb = COLOR_TEXT
            
            grade_str = ""
            if edu.get('cgpa'):
                grade_str = f" (Grade/CGPA: {edu.get('cgpa')})"
            elif edu.get('percentage'):
                grade_str = f" (Percentage: {edu.get('percentage')}%)"
                
            grade = p.add_run(grade_str)
            grade.font.name = 'Arial'
            grade.font.size = Pt(10)
            grade.font.color.rgb = COLOR_TEXT
            
            p.paragraph_format.space_after = Pt(2)
            
            # Pass year (Right aligned if possible, or just trailing)
            p_sub = doc.add_paragraph()
            p_sub.paragraph_format.space_after = Pt(6)
            year_run = p_sub.add_run(f"Graduation: {edu.get('passing_year')}")
            year_run.font.italic = True
            year_run.font.name = 'Arial'
            year_run.font.size = Pt(9.5)
            year_run.font.color.rgb = COLOR_MUTED
            
            if edu.get('achievements'):
                ach_run = p_sub.add_run(f" | Achievements: {edu.get('achievements')}")
                ach_run.font.name = 'Arial'
                ach_run.font.size = Pt(9.5)
                ach_run.font.color.rgb = COLOR_MUTED

    # 4. Experience
    experience = resume_data.get("experience", [])
    if experience:
        add_section_heading("Work Experience")
        for exp in experience:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            
            pos_run = p.add_run(f"{exp.get('position')} ")
            pos_run.font.bold = True
            pos_run.font.name = 'Arial'
            pos_run.font.size = Pt(10)
            pos_run.font.color.rgb = COLOR_TEXT
            
            at_run = p.add_run(f"at {exp.get('company')}")
            at_run.font.name = 'Arial'
            at_run.font.size = Pt(10)
            at_run.font.color.rgb = COLOR_TEXT
            
            dur_p = doc.add_paragraph()
            dur_p.paragraph_format.space_after = Pt(4)
            dur_run = dur_p.add_run(exp.get('duration', ''))
            dur_run.font.italic = True
            dur_run.font.name = 'Arial'
            dur_run.font.size = Pt(9)
            dur_run.font.color.rgb = COLOR_MUTED
            
            desc_p = doc.add_paragraph()
            desc_p.paragraph_format.space_after = Pt(8)
            desc_run = desc_p.add_run(exp.get('description', ''))
            desc_run.font.name = 'Arial'
            desc_run.font.size = Pt(9.5)
            desc_run.font.color.rgb = COLOR_TEXT

    # 5. Projects
    projects = resume_data.get("projects", [])
    if projects:
        add_section_heading("Projects")
        for proj in projects:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            
            name_run = p.add_run(proj.get('name'))
            name_run.font.bold = True
            name_run.font.name = 'Arial'
            name_run.font.size = Pt(10)
            name_run.font.color.rgb = COLOR_TEXT
            
            tech = proj.get('tech_stack')
            if tech:
                tech_run = p.add_run(f" (Tech: {tech})")
                tech_run.font.name = 'Arial'
                tech_run.font.size = Pt(9.5)
                tech_run.font.color.rgb = COLOR_MUTED
                
            dur_p = doc.add_paragraph()
            dur_p.paragraph_format.space_after = Pt(4)
            dur_run = dur_p.add_run(proj.get('duration', ''))
            dur_run.font.italic = True
            dur_run.font.name = 'Arial'
            dur_run.font.size = Pt(9)
            dur_run.font.color.rgb = COLOR_MUTED
            
            desc_p = doc.add_paragraph()
            desc_p.paragraph_format.space_after = Pt(8)
            desc_run = desc_p.add_run(proj.get('description', ''))
            desc_run.font.name = 'Arial'
            desc_run.font.size = Pt(9.5)
            desc_run.font.color.rgb = COLOR_TEXT

    # 6. Technical Skills
    skills = resume_data.get("skills", [])
    if skills:
        add_section_heading("Technical Skills")
        
        # Categorize
        cats = {}
        for s in skills:
            cat = s.get("category", "General")
            cats.setdefault(cat, []).append(f"{s.get('name')} (Lvl {s.get('level')})")
            
        for cat, sk_list in cats.items():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            cat_run = p.add_run(f"{cat}: ")
            cat_run.font.bold = True
            cat_run.font.name = 'Arial'
            cat_run.font.size = Pt(9.5)
            cat_run.font.color.rgb = COLOR_TEXT
            
            items_run = p.add_run(", ".join(sk_list))
            items_run.font.name = 'Arial'
            items_run.font.size = Pt(9.5)
            items_run.font.color.rgb = COLOR_TEXT

    # 7. Certifications
    certificates = resume_data.get("certificates", [])
    if certificates:
        add_section_heading("Certifications")
        for cert in certificates:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            
            name_run = p.add_run(f"• {cert.get('name')} ")
            name_run.font.bold = True
            name_run.font.name = 'Arial'
            name_run.font.size = Pt(9.5)
            name_run.font.color.rgb = COLOR_TEXT
            
            org_run = p.add_run(f"— issued by {cert.get('organization')}")
            org_run.font.name = 'Arial'
            org_run.font.size = Pt(9.5)
            org_run.font.color.rgb = COLOR_TEXT
            
            date = cert.get('issue_date')
            if date:
                date_run = p.add_run(f" ({date})")
                date_run.font.italic = True
                date_run.font.name = 'Arial'
                date_run.font.size = Pt(9)
                date_run.font.color.rgb = COLOR_MUTED

    # 8. Achievements & Extracurriculars
    ach_list = master.get("achievements_list") or resume_data.get("achievements")
    if ach_list:
        ach_json = {}
        if isinstance(ach_list, str):
            try:
                ach_json = json.loads(ach_list)
            except Exception:
                pass
        elif isinstance(ach_list, dict):
            ach_json = ach_list
            
        if ach_json:
            has_ach = False
            for k, v in ach_json.items():
                if v:
                    has_ach = True
                    break
            
            if has_ach:
                add_section_heading("Achievements & Extra-Curriculars")
                for key, val in ach_json.items():
                    if val:
                        p = doc.add_paragraph()
                        p.paragraph_format.space_after = Pt(4)
                        k_run = p.add_run(f"{key.replace('_', ' ').capitalize()}: ")
                        k_run.font.bold = True
                        k_run.font.name = 'Arial'
                        k_run.font.size = Pt(9.5)
                        k_run.font.color.rgb = COLOR_TEXT
                        
                        v_run = p.add_run(val)
                        v_run.font.name = 'Arial'
                        v_run.font.size = Pt(9.5)
                        v_run.font.color.rgb = COLOR_TEXT

    # Save to buffer
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
