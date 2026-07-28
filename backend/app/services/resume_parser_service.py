import io
import urllib.request
import pypdf
import docx
import logging
from typing import Dict, Any

logger = logging.getLogger("resume_parser_service")

def extract_pdf_text(file_bytes: bytes) -> Dict[str, Any]:
    """
    Extracts text from a PDF file using pypdf.
    """
    pdf_file = io.BytesIO(file_bytes)
    reader = pypdf.PdfReader(pdf_file)
    text = ""
    pages = len(reader.pages)
    
    for page in reader.pages:
        extracted = page.extract_text()
        if extracted:
            text += extracted + "\n"
            
    word_count = len(text.split())
    return {
        "text": text.strip(),
        "pages": pages,
        "word_count": word_count
    }

def extract_docx_text(file_bytes: bytes) -> Dict[str, Any]:
    """
    Extracts text from a DOCX file using python-docx.
    """
    docx_file = io.BytesIO(file_bytes)
    doc = docx.Document(docx_file)
    text = ""
    
    for para in doc.paragraphs:
        if para.text:
            text += para.text + "\n"
            
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
            text += "\n"
            
    word_count = len(text.split())
    return {
        "text": text.strip(),
        "pages": 1,  # DOCX doesn't have native page count without rendering
        "word_count": word_count
    }

def extract_doc_text(file_bytes: bytes) -> Dict[str, Any]:
    """
    Fallback extractor for legacy .doc files.
    """
    # Try decoding as raw text as a fallback or extract ASCII strings
    try:
        text = file_bytes.decode('utf-8', errors='ignore')
    except Exception:
        text = ""
    
    # Filter printable ASCII/Unicode characters
    clean_text = "".join(ch for ch in text if ch.isprintable() or ch in "\n\r\t ")
    word_count = len(clean_text.split())
    
    return {
        "text": clean_text.strip(),
        "pages": 1,
        "word_count": word_count
    }

async def extract_resume_text(url: str, filename: str) -> Dict[str, Any]:
    """
    Downloads file from URL and extracts text based on file extension.
    """
    # 1. Download file content
    try:
        # User standard urllib to download
        req = urllib.request.Request(
            url, 
            headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
        )
        with urllib.request.urlopen(req, timeout=15) as response:
            file_bytes = response.read()
    except Exception as e:
        logger.error(f"Failed to download file from Cloudinary URL: {str(e)}")
        raise RuntimeError(f"Download from Cloudinary Failed: {str(e)}")

    # 2. Detect extension
    ext = filename.split(".")[-1].lower() if "." in filename else ""
    
    # 3. Extract text
    if ext == "pdf":
        return extract_pdf_text(file_bytes)
    elif ext == "docx":
        return extract_docx_text(file_bytes)
    elif ext == "doc":
        return extract_doc_text(file_bytes)
    else:
        raise ValueError("Unsupported File Type")
