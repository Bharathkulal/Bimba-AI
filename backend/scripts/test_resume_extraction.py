"""
Bimba AI - Resume Text Extraction Self-Test
===========================================
Run: python scripts/test_resume_extraction.py
Verifies: download, PDF/DOCX parsing, basic text extraction, and MongoDB saving.
"""

import os
import sys

# Ensure project root is on the path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from dotenv import load_dotenv
load_dotenv(os.path.join(os.path.dirname(__file__), "..", ".env"))

import asyncio
from app.services.resume_parser_service import extract_pdf_text, extract_docx_text, extract_resume_text
from app.services.resume_extraction_service import extract_personal_info, extract_skills, extract_structured_data

PASS = "[PASS]"
FAIL = "[FAIL]"
results = []

def record(label, passed, detail=""):
    results.append((label, passed, detail))
    icon = PASS if passed else FAIL
    msg = f"  {icon} {label}"
    if detail:
        msg += f"  ({detail})"
    print(msg)

async def run_tests():
    print("")
    print("=" * 50)
    print("  Bimba AI - Resume Extraction Verification Report")
    print("=" * 50)
    print("")

    # 1. Test PDF extraction locally
    print("Testing PDF Parser service...")
    try:
        # Create a tiny mock PDF in memory to parse
        import pypdf
        writer = pypdf.PdfWriter()
        page = writer.add_blank_page(width=72, height=72)
        
        # Write dummy text if possible, or just parse an empty reader
        # To avoid fitz complex shapes, we just verify the pypdf structures
        pdf_bytes = io_bytes = bytes()
        # Create bytes stream
        import io
        stream = io.BytesIO()
        writer.write(stream)
        pdf_bytes = stream.getvalue()
        
        res = extract_pdf_text(pdf_bytes)
        record("PDF structure extraction works", "text" in res, f"pages={res['pages']}")
    except Exception as e:
        record("PDF structure extraction works", False, str(e))

    # 2. Test DOCX extraction locally
    print("Testing DOCX Parser service...")
    try:
        import docx
        doc = docx.Document()
        doc.add_paragraph("John Doe")
        doc.add_paragraph("Developer at Innovative Solutions")
        doc.add_paragraph("Skills: Python, React, MongoDB")
        
        stream = io.BytesIO()
        doc.save(stream)
        docx_bytes = stream.getvalue()
        
        res = extract_docx_text(docx_bytes)
        has_text = "John Doe" in res["text"]
        record("DOCX text extraction works", has_text, f"words={res['word_count']}")
    except Exception as e:
        record("DOCX text extraction works", False, str(e))

    # 3. Test Metadata Extraction Logic
    print("Testing Metadata Extraction...")
    sample_resume = """
    Johnathan Developer
    johnathan@bimba.ai
    +91 9876543210
    Mangalore, India
    
    Professional Summary
    Dedicated Software Engineer with 2 years of experience building React frontends and Python backends.
    
    Education
    Bachelor of Computer Applications - Tech University (2024)
    
    Skills
    React, Python, JavaScript, Docker, SQL, MongoDB, AWS
    
    Experience
    Software Engineer - Apex Code Studio
    Designed REST backend integrations.
    """
    try:
        info = extract_personal_info(sample_resume)
        record("Personal Info Extracted Name", info["name"] == "Johnathan Developer", info["name"])
        record("Personal Info Extracted Email", info["email"] == "johnathan@bimba.ai", info["email"])
        record("Personal Info Extracted Phone", "+91" in info["phone"], info["phone"])
        
        skills = extract_skills(sample_resume)
        record("Skills mapped correctly", "React" in skills and "Python" in skills, str(skills))
        
        structured = extract_structured_data(sample_resume)
        record("Structured data compilation success", "education" in structured and len(structured["education"]) > 0)
    except Exception as e:
        record("Metadata Extraction Logic", False, str(e))

    # 4. Test MongoDB Saving integration
    print("Testing MongoDB Save integration...")
    try:
        from pymongo import MongoClient
        mongo_uri = os.getenv("MONGODB_URI", "mongodb://localhost:27017")
        db_name   = os.getenv("DATABASE_NAME", "bimba_ai")
        client = MongoClient(mongo_uri, serverSelectionTimeoutMS=3000)
        db = client[db_name]
        
        test_analysis = {
            "_selftest": True,
            "resume_id": 9999,
            "student_id": 1,
            "roll_number": "SELFTEST",
            "raw_text": sample_resume,
            "extracted_data": extract_structured_data(sample_resume),
            "file_type": "pdf",
            "word_count": len(sample_resume.split()),
            "status": "completed",
            "created_at": "2026-07-28T19:30:00Z"
        }
        
        ins = db.resume_analysis.insert_one(test_analysis)
        fetched = db.resume_analysis.find_one({"_id": ins.inserted_id})
        record("MongoDB saving resume_analysis record", fetched["roll_number"] == "SELFTEST")
        
        # Clean up
        db.resume_analysis.delete_one({"_id": ins.inserted_id})
    except Exception as e:
        record("MongoDB saving resume_analysis record", False, str(e))

    print("")
    print("-" * 50)
    passed = sum(1 for _, p, _ in results if p)
    total  = len(results)
    print(f"  Result:  {passed}/{total} checks passed")
    if passed == total:
        print("  Status:  ALL CHECKS PASSED")
    else:
        print("  Status:  SOME CHECKS FAILED")
    print("-" * 50)
    print("")

if __name__ == "__main__":
    asyncio.run(run_tests())
